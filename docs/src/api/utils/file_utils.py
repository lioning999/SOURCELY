"""文件处理工具 — PDF/DOCX 文本提取、清洗、缓存。"""

import asyncio
import atexit
import re
import io
import hashlib
import unicodedata
import pdfplumber
from docx import Document
from concurrent.futures import ThreadPoolExecutor
from config import Config
from .logger import get_logger

logger = get_logger(__name__)

# 专用线程池（PDF/DOCX 解析是 CPU 密集型，不能占事件循环）
FILE_PROCESSING_POOL = ThreadPoolExecutor(max_workers=8, thread_name_prefix="file_processing")
atexit.register(FILE_PROCESSING_POOL.shutdown)

# 文件解析缓存（使用文件内容的哈希值作为键）
FILE_CACHE = {}
CACHE_SIZE_LIMIT = 100  # 缓存大小限制

# AI 结构化结果缓存（键 = cleaned_text MD5，值 = structured_text）
STRUCTURE_CACHE = {}
STRUCTURE_CACHE_SIZE_LIMIT = 100

# 计算文件内容的哈希值
def get_file_hash(content):
    """
    计算文件内容的哈希值，用于缓存键
    
    【作用】
    - 为文件内容生成唯一的哈希值，用于缓存键
    - 确保相同内容的文件只需要处理一次
    
    【参数】
    - content: bytes - 文件二进制内容
    
    【返回值】
    - str: 哈希值字符串
    
    【异常】
    - 无
    """
    return hashlib.md5(content).hexdigest()

# 清理缓存，保持缓存大小在限制内
def cleanup_cache():
    """
    清理缓存，当缓存大小超过限制时删除最旧的项
    
    【作用】
    - 保持缓存大小在合理范围内
    - 避免缓存占用过多内存
    
    【参数】
    - 无
    
    【返回值】
    - 无
    
    【异常】
    - 无
    """
    global FILE_CACHE, STRUCTURE_CACHE
    if len(FILE_CACHE) > CACHE_SIZE_LIMIT:
        keys_to_remove = list(FILE_CACHE.keys())[:len(FILE_CACHE) // 2]
        for key in keys_to_remove:
            del FILE_CACHE[key]
        logger.info(f"FILE_CACHE_CLEANED removed={len(keys_to_remove)} remaining={len(FILE_CACHE)}")
    if len(STRUCTURE_CACHE) > STRUCTURE_CACHE_SIZE_LIMIT:
        keys_to_remove = list(STRUCTURE_CACHE.keys())[:len(STRUCTURE_CACHE) // 2]
        for key in keys_to_remove:
            del STRUCTURE_CACHE[key]
        logger.info(f"STRUCTURE_CACHE_CLEANED removed={len(keys_to_remove)} remaining={len(STRUCTURE_CACHE)}")

def get_text_hash(text: str) -> str:
    """计算文本 MD5，用于结构化缓存键"""
    return hashlib.md5(text.encode('utf-8')).hexdigest()

async def _ocr_pdf_with_tesseract(file_bytes):
    """OCR 扫描件 PDF（图片格式），用 Tesseract 提取文本。在线程池内执行。"""

    def sync_ocr():
        import pypdfium2 as pdfium
        import pytesseract
        import shutil

        # systemd 下 PATH 可能不包含 /usr/bin，显式指定
        tesseract_bin = shutil.which('tesseract') or '/usr/bin/tesseract'
        pytesseract.pytesseract.tesseract_cmd = tesseract_bin
        pdf = pdfium.PdfDocument(file_bytes)
        text_parts = []
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=2)
            img = bitmap.to_pil()
            page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            bitmap.close()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
        pdf.close()
        return '\n'.join(text_parts)

    return await asyncio.get_event_loop().run_in_executor(
        FILE_PROCESSING_POOL, sync_ocr,
    )


async def extract_text_from_pdf(file_bytes):
    """从 PDF 提取文本。先 pdfplumber，文字为空时降级 Tesseract OCR。"""

    def sync_extract():
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    logger.info(f"PDF_EXTRACT_START file_size={len(file_bytes)}")
    try:
        text = await asyncio.get_event_loop().run_in_executor(
            FILE_PROCESSING_POOL, sync_extract,
        )
    except Exception as e:
        logger.error(f"PDF_EXTRACT_ERROR error={str(e)}")
        raise

    if not text or not text.strip():
        logger.info("PDF_PLUMBER_EMPTY falling back to Tesseract OCR")
        try:
            text = await _ocr_pdf_with_tesseract(file_bytes)
            logger.info(f"PDF_TESSERACT_DONE text_len={len(text)}")
        except Exception as e:
            logger.error(f"PDF_TESSERACT_ERROR error={str(e)}")
            # 不抛——返回空文本，让上层校验给出友好提示

    return text

async def extract_text_from_docx(file_bytes):
    """
    从Word(.docx)二进制数据中提取文本
    
    【作用】
    - 从Word文件中提取文本内容
    - 使用线程池执行同步操作，避免阻塞事件循环
    
    【参数】
    - file_bytes: bytes - 文件的二进制内容
    
    【返回值】
    - str: 提取的文本字符串
    
    【异常】
    - Exception: Word文档解析失败时抛出
    """
    def sync_extract():
        text = ""
        try:
            # 使用python-docx打开文档，同样需要BytesIO
            doc = Document(io.BytesIO(file_bytes))
            # 提取所有段落的文本
            for para in doc.paragraphs:
                text += para.text + "\n"
            logger.info(f"DOCX_EXTRACT_SUCCESS paragraphs={len(doc.paragraphs)}")
        except Exception as e:
            logger.error(f"DOCX_EXTRACT_ERROR error={str(e)}")
            raise
        return text
    
    # 使用专用线程池执行同步操作
    return await asyncio.get_event_loop().run_in_executor(FILE_PROCESSING_POOL, sync_extract)

async def extract_resume_text(content: bytes, filename: str) -> str:
    """从文件字节提取文本。PDF/DOCX 用线程池解析，TXT 直接 decode。无 FastAPI 依赖。"""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext == 'pdf':
        text = await extract_text_from_pdf(content)
    elif ext == 'docx':
        text = await extract_text_from_docx(content)
    elif ext == 'txt':
        text = content.decode('utf-8', errors='replace')
    else:
        raise ValueError(f"不支持的文件格式: .{ext}")

    return await clean_resume_text(text)


# ── CJK 字符标准化 ──────────────────────────────────────
# PDF 提取时可能把标准汉字存成康熙部首（U+2F00–U+2FDF）或
# CJK 兼容区码点（U+F900–U+FAFF）。屏幕看起来一样，编码不同。
# 这套映射表把常见变体转回标准汉字，避免 AI 浪费 token 修正编码。

_CJK_NORMALIZE_MAP = {
    # 康熙部首 → 标准汉字（pdfplumber 常见问题字符）
    '⻜': '飞',   # ⻜ → 飞
    '⻛': '行',   # ⻛ → 行
    '⻓': '手',   # ⻓ → 手
    '⻘': '人',   # ⻘ → 人
    '⻖': '门',   # ⻖ → 门
    '⻙': '比',   # ⻙ → 比
    '⼴': '广',   # ⼴ → 广
    '⼚': '大',   # ⼚ → 大
    '⼈': '方',   # ⼈ → 方
    '⼍': '力',   # ⼍ → 力
    '⼩': '官',   # ⼩ → 官
    # CJK 兼容区常见映射
    '兀': '連',   # 兀 → 連
    '嗀': '藏',   # 嗀 → 藏
}

def normalize_cjk(text: str) -> tuple[str, int]:
    """将 PDF 提取的康熙部首/兼容区字符转回标准汉字。

    先用 NFKC 标准化（处理大部分兼容字符），再补自定义映射表
    （NFKC 对康熙部首映射不完整）。

    Returns:
        (标准化后的文本, 修正的字符数)
    """
    # Step 1: NFKC 标准化
    nfkc = unicodedata.normalize('NFKC', text)

    # Step 2: 补自定义映射
    fixed = 0
    result = []
    for ch in nfkc:
        mapped = _CJK_NORMALIZE_MAP.get(ch)
        if mapped is not None:
            result.append(mapped)
            fixed += 1
        else:
            result.append(ch)
    return ''.join(result), fixed


async def clean_resume_text(text):
    """
    清洗简历文本：去除多余空行、特殊字符，限制长度
    
    【作用】
    - 清洗简历文本，去除多余空行和特殊字符
    - 限制文本长度，确保符合AI输入要求
    
    【参数】
    - text: str - 原始文本
    
    【返回值】
    - str: 清洗后的文本
    
    【异常】
    - 无
    """
    # 1. 按行分割，去除每行首尾空格，过滤掉空行
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    # 2. 重新组合成文本，保留换行分隔
    cleaned = "\n".join(lines)
    
    # 3. 去除不可打印控制字符，保留所有可打印字符（含邮箱/URL/特殊符号）
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)

    # 3.5 CJK 字符标准化：康熙部首/兼容区码点 → 标准汉字
    cleaned, cjk_fixed = normalize_cjk(cleaned)
    if cjk_fixed > 0:
        logger.info(f"CJK_NORMALIZED fixed={cjk_fixed}")

    # 4. 限制长度（AI输入限制，从配置读取）
    max_length = Config.RESUME_TEXT_MAX_LENGTH
    original_length = len(cleaned)
    if len(cleaned) > max_length:
        cleaned = cleaned[:max_length]
        logger.warning(f"TEXT_TRUNCATED original_length={original_length} truncated_to={max_length}")
    
    logger.info(f"TEXT_CLEANED original_length={original_length} cleaned_length={len(cleaned)}")
    return cleaned

async def parse_resume(file, content=None):
    """
    主函数：根据上传的文件对象，解析并清洗文本
    
    【作用】
    - 统一处理文件上传、解析和清洗流程
    - 实现文件解析结果的缓存，提高处理效率
    
    【参数】
    - file: UploadFile - FastAPI的UploadFile对象（来自请求）
    - content: bytes, optional - 文件的二进制内容，如果提供则直接使用，否则从file对象中读取
    
    【返回值】
    - str: 清洗后的文本字符串
    
    【异常】
    - ValueError: 不支持的文件格式
    - Exception: 文件处理失败时抛出
    """
    # 读取文件内容（二进制）
    if content is None:
        content = await file.read()
    
    file_size = len(content)
    logger.info(f"FILE_UPLOAD_START file_size={file_size} filename={file.filename}")
    
    # 计算文件哈希值，用于缓存
    file_hash = get_file_hash(content)
    
    # 检查缓存
    if file_hash in FILE_CACHE:
        logger.info(f"FILE_CACHE_HIT hash={file_hash[:10]}...")
        return FILE_CACHE[file_hash]
    
    # 获取文件名并提取扩展名
    filename = file.filename
    ext = filename.split('.')[-1].lower()
    
    try:
        if ext == 'pdf':
            text = await extract_text_from_pdf(content)
        elif ext == 'docx':
            text = await extract_text_from_docx(content)
        else:
            logger.error(f"UNSUPPORTED_FILE_TYPE ext={ext} filename={filename}")
            raise ValueError("不支持的文件格式，请上传PDF或Word文档（.docx）")
        
        # 清洗文本
        cleaned = await clean_resume_text(text)
        
        # 缓存结果
        FILE_CACHE[file_hash] = cleaned
        logger.info(f"FILE_CACHE_STORED hash={file_hash[:10]}... size={len(cleaned)}")
        
        # 清理缓存（如果需要）
        cleanup_cache()
        
        logger.info(f"FILE_PROCESSING_COMPLETE filename={filename} status=success")
        return cleaned
    except Exception as e:
        logger.error(f"FILE_PROCESSING_ERROR filename={filename} error={str(e)}")
        raise